"""
HSK 智能文章生成器 - Streamlit 网页版
使用 Streamlit 框架快速部署
"""

import streamlit as st
import requests
import re
from xpinyin import Pinyin
import jieba
import os
import tempfile
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ================== 页面配置 ==================
st.set_page_config(
    page_title="🎓 HSK 智能文章生成器",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ================== API配置 ==================
API_KEY = "sk-cad3c95291134e868ca15ade100c1033"

# ================== 初始化 ==================
pinyin_tool = Pinyin()

# 初始化 session state
if 'generated_articles' not in st.session_state:
    st.session_state.generated_articles = []
if 'mp3_cache' not in st.session_state:
    st.session_state.mp3_cache = {}
if 'word_cache' not in st.session_state:
    st.session_state.word_cache = {}
if 'generation_complete' not in st.session_state:
    st.session_state.generation_complete = False
if 'current_level' not in st.session_state:
    st.session_state.current_level = None
if 'current_words' not in st.session_state:
    st.session_state.current_words = None
if 'include_mp3_option' not in st.session_state:
    st.session_state.include_mp3_option = False

# HSK标准信息
HSK_INFO = {
    "HSK1": {
        "词汇量": "150词",
        "语法": "主谓宾、的、很、吗、呢、这那、想、可以、会（能力）、在+地点、有/没有",
        "话题": "问候、自我介绍、家庭成员、数字、时间、简单购物、基本爱好"
    },
    "HSK2": {
        "词汇量": "300词（累计）",
        "语法": "了（完成）、过（经历）、在+动词、着、比、越…越…、刚、正在、从…到…、虽然…但是…、因为…所以…",
        "话题": "日常生活、简单工作、基本旅行、天气、身体状况"
    },
    "HSK3": {
        "词汇量": "600词（累计）",
        "语法": "除了…以外、一边…一边…、只要…就…、尽管…还是…、反而、难道、即使…也…、把字句、被字句",
        "话题": "学校生活、工作场景、旅行经历、健康养生、节日庆典"
    },
    "HSK4": {
        "词汇量": "1200词（累计）",
        "语法": "所谓、毕竟、简直、竟然、看来、显然、幸亏、难免、至于、从而、由此可见、相比之下、动不动就、连…都…",
        "话题": "职场、社会现象、文化差异、科技生活、环境问题"
    },
    "HSK5": {
        "词汇量": "2500词（累计）",
        "语法": "成语、四字词语、以至于、之所以…是因为…、宁可…也不…、无论…都…、与其…不如…",
        "话题": "深度社会议题、传统文化、科技发展、哲学思考、职业规划"
    }
}

# ================== API调用函数 ==================
def qwen3_generate(prompt):
    """调用通义千问生成文本"""
    url = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "qwen-max",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.8,
        "top_p": 0.9,
        "max_tokens": 1500
    }
    
    resp = requests.post(url, headers=headers, json=data, timeout=90)
    if resp.status_code == 200:
        result = resp.json()
        return result['choices'][0]['message']['content']
    else:
        raise Exception(f"API错误: {resp.status_code}, {resp.text}")

def qwen3_tts(text, save_path):
    """调用qwen3-tts-flash生成语音（墨讲师，0.8倍速）"""
    url = "https://dashscope.aliyuncs.com/api/v1/services/aigc/multimodal-generation/generation"
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "qwen3-tts-flash",
        "input": {
            "text": text,
            "voice": "Elias",
            "language_type": "Chinese",
            "rate": 0.8
        }
    }

    resp = requests.post(url, headers=headers, json=data, timeout=60)
    if resp.status_code != 200:
        raise Exception(f"TTS API错误: {resp.status_code}, {resp.text}")

    result = resp.json()

    # 兼容返回格式：可能是直接返回音频 URL，也可能返回 base64 数据
    audio_bytes = None

    # 处理 output.audio.url
    try:
        audio_info = result.get('output', {}).get('audio', {})
    except Exception:
        audio_info = {}

    # 优先处理 base64 data 字段
    audio_data_b64 = audio_info.get('data') or audio_info.get('b64') or audio_info.get('base64')
    if audio_data_b64:
        import base64
        try:
            audio_bytes = base64.b64decode(audio_data_b64)
        except Exception as e:
            raise Exception(f"无法解析 TTS 返回的 base64 音频数据: {e}")

    # 否则如果有 URL，则去请求该 URL
    if audio_bytes is None:
        audio_url = audio_info.get('url')
        if not audio_url:
            # 有些返回把 audio 放在不同层级，尝试直接查找
            if isinstance(result.get('output'), dict):
                for v in result['output'].values():
                    if isinstance(v, dict) and ('url' in v or 'data' in v):
                        audio_url = v.get('url')
                        audio_data_b64 = audio_data_b64 or v.get('data')
                        break

        if audio_data_b64 and audio_bytes is None:
            import base64
            try:
                audio_bytes = base64.b64decode(audio_data_b64)
            except Exception as e:
                raise Exception(f"无法解析 TTS 返回的 base64 音频数据: {e}")

        if audio_bytes is None and audio_url:
            audio_resp = requests.get(audio_url, timeout=60)
            if audio_resp.status_code != 200:
                raise Exception(f"无法下载音频文件: {audio_resp.status_code}, {audio_resp.text}")
            audio_bytes = audio_resp.content

    if not audio_bytes:
        raise Exception("TTS 未返回可用的音频数据")

    # 写入文件（API 返回的是 WAV 格式）
    with open(save_path, "wb") as f:
        f.write(audio_bytes)
    
    print(f"[调试] 音频文件已保存: {save_path}, 大小: {len(audio_bytes)} 字节")

# ================== Word文档生成函数 ==================
def add_pinyin_to_word(doc, text, red_words):
    """将文本添加到 Word 文档，拼音在上，汉字在下"""
    chars = list(text)
    max_chars_per_line = 10
    
    i = 0
    while i < len(chars):
        line_chars = []
        for _ in range(max_chars_per_line):
            if i < len(chars):
                line_chars.append(chars[i])
                i += 1
            else:
                break
        
        if not line_chars:
            continue
        
        # 创建表格
        table = doc.add_table(rows=2, cols=len(line_chars))
        table.style = 'Table Grid'
        table.allow_autofit = False
        table.width = Inches(6.5)
        
        for j, char in enumerate(line_chars):
            table.columns[j].width = Inches(0.55)
            
            # 拼音行
            py_cell = table.cell(0, j)
            if re.match(r'[\u4e00-\u9fff]', char):  # 汉字
                py = pinyin_tool.get_pinyin(char, tone_marks='marks')
                py_cell.text = py
                if py_cell.paragraphs[0].runs:
                    py_run = py_cell.paragraphs[0].runs[0]
                    py_run.font.size = Pt(12)
                    is_red = any(char in word for word in red_words)
                    py_run.font.color.rgb = RGBColor(255, 0, 0) if is_red else RGBColor(0, 0, 255)
            else:
                py_cell.text = ""
            
            py_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 汉字行
            text_cell = table.cell(1, j)
            text_cell.text = char
            p = text_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if p.runs:
                run = p.runs[0]
                run.font.name = 'SimSun'
                run.font.size = Pt(16)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                is_red_word = any(char in word for word in red_words)
                if is_red_word:
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.bold = True
        
        doc.add_paragraph()

# ================== 侧边栏配置 ==================
st.sidebar.markdown("## ⚙️ 生成参数")

with st.sidebar:
    st.markdown("### 📝 输入生词")
    words_input = st.text_area(
        "生词列表",
        placeholder="用空格或回车分隔生词\n例如：学生 学校 上课",
        height=100,
        label_visibility="collapsed"
    )
    
    st.markdown("### 🎯 HSK 级别")
    level = st.selectbox(
        "选择级别",
        ["HSK1", "HSK2", "HSK3", "HSK4", "HSK5"],
        index=2,
        label_visibility="collapsed"
    )
    
    st.markdown("### 📄 生成篇数")
    num_articles = st.slider(
        "篇数",
        1, 10, 3,
        label_visibility="collapsed"
    )
    
    st.markdown("### 📏 每篇字数")
    text_length = st.selectbox(
        "字数范围",
        ["1-50字", "50-100字", "100-200字", "200字以上"],
        index=2,
        label_visibility="collapsed"
    )
    
    st.markdown("### 🔊 生成选项")
    include_mp3 = st.checkbox("生成 MP3 朗读（墨讲师）", value=False)
    
    st.divider()
    
    # HSK 信息显示
    st.markdown(f"### 📚 {level} 级别标准")
    hsk_info = HSK_INFO[level]
    st.markdown(f"""
    **词汇量**：{hsk_info['词汇量']}
    
    **语法要点**：
    {hsk_info['语法'][:100]}...
    
    **话题范围**：
    {hsk_info['话题']}
    """)

# ================== 主页面 ==================
st.title("🎓 HSK 智能文章生成器")
st.markdown("基于通义千问 AI · 专业级 HSK 学习材料生成工具")

# 生成按钮
if st.button("🚀 开始生成", key="generate_btn", use_container_width=True):
    # 验证输入
    if not words_input.strip():
        st.error("❌ 请先输入生词！")
    else:
        # 处理生词
        words = [w.strip() for w in words_input.replace("\n", " ").split() if w.strip()]
        
        # 保存当前配置到 session_state
        st.session_state.current_level = level
        st.session_state.current_words = words
        st.session_state.include_mp3_option = include_mp3
        
        # 创建进度条
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        all_articles = []
        
        try:
            for i in range(num_articles):
                # 更新进度
                progress = (i + 1) / num_articles
                progress_bar.progress(progress)
                status_text.text(f"⏳ 正在生成第 {i+1}/{num_articles} 篇文章...")
                
                # 生成提示词
                length_map = {
                    "1-50字": "40字左右",
                    "50-100字": "80字左右",
                    "100-200字": "150字左右",
                    "200字以上": "400字左右"
                }
                target_length = length_map[text_length]
                
                hsk_info = HSK_INFO[level]
                story_types = {
                    "HSK1": "日常生活小故事（如：买水果、问路、介绍家人）",
                    "HSK2": "简单生活故事（如：去旅行、看医生、过生日）",
                    "HSK3": "有趣的经历故事（如：学习中文、中国节日、健康生活）",
                    "HSK4": "文化故事（如：中国传统文化、历史人物、名胜古迹）",
                    "HSK5": "深度文化故事（如：中国历史典故、文化传统、社会现象）"
                }
                
                prompt = f"""你是一位经验丰富的对外汉语教师，擅长创作引人入胜的中文学习材料。

【HSK{level[-1]}级别标准】
- 词汇量：{hsk_info['词汇量']}
- 允许使用的语法：{hsk_info['语法']}
- 话题范围：{hsk_info['话题']}

【创作要求】
请创作一篇严格符合HSK{level[-1]}水平的趣味短文：

1. 📖 内容形式：写成{story_types[level]}，要有完整的故事情节或清晰的逻辑结构
2. 📚 词汇要求：只能使用HSK1到HSK{level[-1]}范围内的词汇
3. ✏️ 语法要求：只能使用上述允许的语法结构
4. 🎯 生词融入：必须自然地包含以下所有生词，每个生词至少出现1-2次：{' '.join(words)}
5. 📏 字数要求：{target_length}
6. 🎨 写作风格：语言生动、有画面感、逻辑清晰、易于理解

【标点符号要求】
- 必须使用正确的中文标点符号
- 句号（。）、逗号（，）、问号（？）、感叹号（！）
- 每句话必须有标点符号

【输出格式】
只输出纯汉字文章（必须带标点符号），不要拼音、不要编号、不要任何额外标记。

请创作一篇让学生读完后既学到知识又感到有趣的短文。"""
                
                # 调用 API
                article = qwen3_generate(prompt)
                
                # 清理文章
                article = re.sub(r"[*#\[\]【】\n\r\t]", "", article)
                article = article.strip()
                
                # 如果没有标点，自动添加
                if not re.search(r'[。！？]', article):
                    chars = list(article)
                    result = []
                    count = 0
                    for idx, char in enumerate(chars):
                        result.append(char)
                        count += 1
                        if count >= 10 and count <= 15 and idx < len(chars) - 1:
                            result.append('。')
                            count = 0
                    if result and result[-1] not in '。！？':
                        result.append('。')
                    article = ''.join(result)
                
                # 字数控制
                max_length = 450 if text_length == "200字以上" else 250
                if len(article) > max_length:
                    article = article[:max_length]
                
                all_articles.append(article)
            
            # 清除进度条
            progress_bar.empty()
            status_text.empty()
            
            # 缓存生成的文章和标志
            st.session_state.generated_articles = all_articles
            st.session_state.generation_complete = True
            
            # 显示成功消息
            st.success(f"✅ 成功生成 {num_articles} 篇文章！")
            st.divider()
            
        except Exception as e:
            st.error(f"❌ 生成失败：{str(e)}")
            st.info("可能原因：\n- API 密钥无效\n- 网络连接问题\n- API 服务不可用")

# ================== 显示生成结果（独立于生成按钮） ==================
# 如果有已生成的文章，显示它们（即使没有再次点击生成按钮）
if st.session_state.generation_complete and st.session_state.generated_articles:
    st.divider()
    st.markdown("### 📖 生成的文章")
    
    level = st.session_state.current_level
    words = st.session_state.current_words
    include_mp3 = st.session_state.include_mp3_option
    
    # 显示生成的文章
    for i, article in enumerate(st.session_state.generated_articles, 1):
        with st.expander(f"📖 文章 {i}", expanded=(i==1)):
            st.markdown(f"**字数**: {len(article)} 字")
            
            # 显示原文
            st.markdown("**原文：**")
            st.info(article)
            
            # Word 文档缓存 key
            word_key = f"word_{level}_{i}"
            
            # 如果缓存中没有 Word 文档，则生成
            if word_key not in st.session_state.word_cache:
                doc = Document()
                doc.add_heading(f"HSK{level} 第 {i} 篇", level=1)
                doc.add_paragraph()
                add_pinyin_to_word(doc, article, words)
                
                # 转换为字节
                word_bytes = io.BytesIO()
                doc.save(word_bytes)
                word_bytes.seek(0)
                
                # 缓存 Word 字节
                st.session_state.word_cache[word_key] = word_bytes.getvalue()
            
            # 从缓存取出 Word 数据
            word_data = st.session_state.word_cache[word_key]
            
            # 下载 Word 文档（带拼音标注）
            st.download_button(
                label="📥 下载 Word（带拼音标注）",
                data=word_data,
                file_name=f"HSK{level}_文章{i}_带拼音.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"btn_word_{i}",
                use_container_width=True
            )
            
            # MP3 部分
            if include_mp3:
                st.markdown("---")
                st.markdown("**🎵 MP3 朗读（墨讲师）**")
                
                mp3_key = f"mp3_{level}_{i}"
                
                # 检查缓存中是否已有该 MP3
                if mp3_key not in st.session_state.mp3_cache:
                    # 显示"生成 MP3"按钮
                    if st.button(f"🎬 生成 MP3 {i}", key=f"generate_mp3_{i}", use_container_width=True):
                        try:
                            progress_placeholder = st.empty()
                            progress_placeholder.text(f"🔄 正在生成第 {i} 篇的 MP3 音频...")
                            
                            # 创建临时文件来存储 MP3
                            with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as tmp_file:
                                tmp_path = tmp_file.name
                            
                            qwen3_tts(article, tmp_path)
                            
                            # 读取 MP3 文件
                            with open(tmp_path, 'rb') as mp3_file:
                                mp3_data = mp3_file.read()
                            
                            # 清理临时文件
                            os.unlink(tmp_path)
                            
                            # 缓存 MP3 数据（这会自动刷新页面）
                            st.session_state.mp3_cache[mp3_key] = mp3_data
                            progress_placeholder.empty()
                            st.rerun()
                            
                        except Exception as e:
                            st.error(f"❌ MP3 生成失败：{str(e)}")
                            st.info("可能原因：API 密钥无效、网络连接问题或 API 服务不可用")
                else:
                    # 显示已缓存的 MP3/WAV
                    audio_data = st.session_state.mp3_cache[mp3_key]
                    
                    # 显示音频播放器
                    st.audio(audio_data, format='audio/wav')
                    
                    # 显示下载按钮（不会导致内容消失）
                    st.download_button(
                        label=f"📥 下载音频 {i}",
                        data=audio_data,
                        file_name=f"HSK{level}_第{i}篇_墨讲师朗读.wav",
                        mime="audio/wav",
                        key=f"btn_mp3_{i}",
                        use_container_width=True
                    )


# ================== 底部信息 ==================
st.divider()
st.markdown("""
---
### 📌 使用说明

1. **输入生词** - 在左侧输入要学习的词汇，支持空格或回车分隔
2. **选择级别** - 根据学习阶段选择 HSK 等级（1-5）
3. **配置参数** - 调整篇数和字数
4. **点击生成** - 等待 AI 生成个性化学习材料
5. **下载使用** - 支持 Word、TXT、MP3 等多种格式

### 🎯 功能特点

✨ **AI 智能生成** - 基于通义千问大模型
📖 **多级别支持** - 从 HSK1 到 HSK5
🔤 **拼音标注** - 一个字一格，清晰易读
🔴 **生词标记** - 自动标红生词，强化学习
📥 **Word 下载** - 带拼音标注的 Word 文档
🎵 **MP3 朗读** - 墨讲师专业朗读，0.8倍速

### 💡 技术栈

- **前端框架** - Streamlit
- **后端 AI** - 通义千问 3-MAX
- **拼音库** - xpinyin
- **分词工具** - jieba

---
**版本**: 2.0 | **最后更新**: 2025-12-11
""")
